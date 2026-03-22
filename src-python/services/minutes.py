"""
Minutes services — Markdown/DOCX conversation and legacy payload parsing.
"""

import json
import re
import tempfile
from pathlib import Path

from logger import get_logger

log = get_logger(__name__)


def parse_legacy_minutes_payload(raw: str):
    """Parse legacy JSON-wrapped minutes payload. Returns dict or None."""
    text = (raw or "").strip()
    if not text:
        return None
    if text.startswith("```"):
        lines = text.splitlines()
        if len(lines) >= 2 and lines[0].startswith("```") and lines[-1].strip() == "```":
            text = "\n".join(lines[1:-1]).strip()
    if not (text.startswith("{") and text.endswith("}")):
        return None
    try:
        data = json.loads(text)
    except Exception:
        return None
    if not isinstance(data, dict):
        return None
    keys = {"title", "attendees", "keyPoints", "decisions", "actionItems", "summary"}
    if not any(k in data for k in keys):
        return None
    return data


def legacy_minutes_to_markdown(data: dict, lang: str = "vi") -> str:
    """Convert legacy minutes JSON dict to Markdown string."""
    vi = lang == "vi"
    title = str(data.get("title") or "").strip() or ("Biên bản cuộc họp" if vi else "Meeting Minutes")
    attendees = [str(x).strip() for x in (data.get("attendees") or []) if isinstance(x, str) and str(x).strip()]
    key_points = [str(x).strip() for x in (data.get("keyPoints") or []) if isinstance(x, str) and str(x).strip()]
    decisions = [str(x).strip() for x in (data.get("decisions") or []) if isinstance(x, str) and str(x).strip()]
    summary = str(data.get("summary") or "").strip()
    action_rows = []
    for item in data.get("actionItems") or []:
        if not isinstance(item, dict):
            continue
        task = str(item.get("task") or "").strip()
        if not task:
            continue
        action_rows.append({
            "task": task,
            "assignee": str(item.get("assignee") or "").strip(),
            "deadline": str(item.get("deadline") or "").strip(),
        })

    missing = "Chưa có dữ liệu" if vi else "Missing data"
    unknown = "Chưa rõ" if vi else "TBD"
    blocks = [
        f"# {title}",
        f"## {'Thành phần tham gia' if vi else 'Attendees'}\n" + ("\n".join(f"- {x}" for x in attendees) if attendees else f"- {missing}"),
        f"## {'Nội dung trao đổi chính' if vi else 'Key Discussion'}\n" + ("\n".join(f"- {x}" for x in key_points) if key_points else f"- {missing}"),
        f"## {'Quyết định quan trọng' if vi else 'Key Decisions'}\n" + ("\n".join(f"- {x}" for x in decisions) if decisions else f"- {missing}"),
    ]
    if action_rows:
        action_lines = []
        for idx, row in enumerate(action_rows, start=1):
            action_lines.append(
                f"{idx}. **What:** {row['task']}\n"
                f"   **Who:** {row['assignee'] or unknown}\n"
                f"   **When:** {row['deadline'] or unknown}"
            )
        blocks.append(f"## {'Action items (What - Who - When)' if vi else 'Action Items (What - Who - When)'}\n" + "\n".join(action_lines))
    else:
        blocks.append(f"## {'Action items (What - Who - When)' if vi else 'Action Items (What - Who - When)'}\n- {missing}")

    if summary:
        blocks.append(f"## {'Tóm tắt' if vi else 'Summary'}\n{summary}")
    return "\n\n".join(blocks).strip()


def normalize_minutes_markdown(raw: str, lang: str = "vi") -> str:
    """Normalize minutes content — handles both Markdown and legacy JSON formats."""
    text = (raw or "").strip()
    if not text:
        return ""
    legacy = parse_legacy_minutes_payload(text)
    if legacy:
        return legacy_minutes_to_markdown(legacy, lang)
    return text


def markdown_to_docx(markdown_text: str, out_path: Path) -> None:
    """Convert Markdown text to a .docx file."""
    from docx import Document

    def _strip_inline(md_line: str) -> str:
        line = re.sub(r"`([^`]+)`", r"\1", md_line)
        line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        line = re.sub(r"\*([^*]+)\*", r"\1", line)
        return line.strip()

    doc = Document()
    for raw in markdown_text.replace("\r", "").split("\n"):
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue
        h = re.match(r"^(#{1,3})\s+(.+)$", line)
        if h:
            level = min(3, len(h.group(1)))
            doc.add_heading(_strip_inline(h.group(2)), level=level)
            continue
        ul = re.match(r"^[-*]\s+(.+)$", line)
        if ul:
            doc.add_paragraph(_strip_inline(ul.group(1)), style="List Bullet")
            continue
        ol = re.match(r"^\d+\.\s+(.+)$", line)
        if ol:
            doc.add_paragraph(_strip_inline(ol.group(1)), style="List Number")
            continue
        doc.add_paragraph(_strip_inline(line))
    doc.save(out_path)
