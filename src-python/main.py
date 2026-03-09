"""
VoiceScribe Python Sidecar — FastAPI server
Runs as a local HTTP server spawned by Tauri.
Handles: STT (Groq), Speaker Diarization (Pyannote), Translation, Summarization
"""

# ── Windows PyInstaller: must be first ──
import multiprocessing
import sys
import os
multiprocessing.freeze_support()

import asyncio
import json
import subprocess
import tempfile
import time
from contextlib import asynccontextmanager
from pathlib import Path
from uuid import uuid4

import uvicorn
from fastapi import FastAPI, UploadFile, File, Request, WebSocket, WebSocketDisconnect
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse, FileResponse, Response
from starlette.background import BackgroundTask

from db import Database
from diarize import SpeakerDiarizer
from i18n import t
from stt import transcribe_nvidia, NvidiaStreamingSTT, HALLUCINATION_PATTERNS, get_language_code
import collections
import logging
import threading
import queue as _queue


# ─── Log Ring Buffer ───
MAX_LOG_LINES = 500
_log_buffer: collections.deque = collections.deque(maxlen=MAX_LOG_LINES)
_log_subscribers: list[_queue.Queue] = []
_log_lock = threading.Lock()


class _LogInterceptor:
    """Intercepts writes to stdout/stderr and copies them to the log buffer."""
    def __init__(self, original):
        self._original = original

    def write(self, msg):
        self._original.write(msg)
        if msg and msg.strip():
            ts = time.strftime("%H:%M:%S")
            line = f"[{ts}] {msg.rstrip()}"
            with _log_lock:
                _log_buffer.append(line)
                for q in list(_log_subscribers):
                    try:
                        q.put_nowait(line)
                    except _queue.Full:
                        pass

    def flush(self):
        self._original.flush()

    def fileno(self):
        return self._original.fileno()

    def isatty(self):
        return False


sys.stdout = _LogInterceptor(sys.stdout)
sys.stderr = _LogInterceptor(sys.stderr)

# ─── Globals ───
db = Database()
diarizer = SpeakerDiarizer()


def _voicescribe_data_dir() -> Path:
    env_dir = os.getenv("VOICESCRIBE_DATA")
    if env_dir:
        return Path(env_dir)
    db_path = getattr(db, "_db_path", None)
    if db_path:
        return Path(db_path).parent
    return Path.home() / ".voicescribe"


def _safe_unlink(path: str):
    try:
        p = Path(path)
        if p.exists() and p.is_file():
            p.unlink()
    except Exception:
        pass


def _transcode_audio_for_export(source: Path, fmt: str) -> Path:
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=f".{fmt}")
    tmp_path = Path(tmp.name)
    tmp.close()

    # Internal system-audio archive may be raw PCM16 (16kHz mono)
    if source.suffix.lower() == ".pcm":
        input_args = ["-f", "s16le", "-ar", "16000", "-ac", "1", "-i", str(source)]
    else:
        input_args = ["-i", str(source)]

    if fmt == "wav":
        cmd = ["ffmpeg", "-y", *input_args, "-vn", "-acodec", "pcm_s16le", "-ar", "16000", "-ac", "1", str(tmp_path)]
    else:  # mp4
        cmd = ["ffmpeg", "-y", *input_args, "-vn", "-acodec", "aac", "-b:a", "192k", str(tmp_path)]

    result = subprocess.run(cmd, capture_output=True, timeout=120)
    if result.returncode != 0:
        _safe_unlink(str(tmp_path))
        stderr = result.stderr.decode(errors="replace")[:300]
        raise RuntimeError(stderr or "ffmpeg convert failed")
    return tmp_path


def _audio_media_type(ext: str) -> str:
    ext_l = ext.lower()
    if ext_l == ".wav":
        return "audio/wav"
    if ext_l == ".mp4":
        return "audio/mp4"
    if ext_l == ".mp3":
        return "audio/mpeg"
    if ext_l == ".m4a":
        return "audio/mp4"
    return "application/octet-stream"


def _parse_legacy_minutes_payload(raw: str):
    text = (raw or "").strip()
    if not text:
        return None
    if text.startswith("```"):
        lines = text.splitlines()
        if len(lines) >= 2 and lines[0].startswith("```") and lines[-1].strip() == "```":
            text = "\n".join(lines[1:-1]).strip()
    if not (text.startswith("{") and text.endswith("}")):
        return None
    try:
        data = json.loads(text)
    except Exception:
        return None
    if not isinstance(data, dict):
        return None
    keys = {"title", "attendees", "keyPoints", "decisions", "actionItems", "summary"}
    if not any(k in data for k in keys):
        return None
    return data


def _legacy_minutes_to_markdown(data: dict, lang: str = "vi") -> str:
    vi = lang == "vi"
    title = str(data.get("title") or "").strip() or ("Biên bản cuộc họp" if vi else "Meeting Minutes")
    attendees = [str(x).strip() for x in (data.get("attendees") or []) if isinstance(x, str) and str(x).strip()]
    key_points = [str(x).strip() for x in (data.get("keyPoints") or []) if isinstance(x, str) and str(x).strip()]
    decisions = [str(x).strip() for x in (data.get("decisions") or []) if isinstance(x, str) and str(x).strip()]
    summary = str(data.get("summary") or "").strip()
    action_rows = []
    for item in data.get("actionItems") or []:
        if not isinstance(item, dict):
            continue
        task = str(item.get("task") or "").strip()
        if not task:
            continue
        action_rows.append({
            "task": task,
            "assignee": str(item.get("assignee") or "").strip(),
            "deadline": str(item.get("deadline") or "").strip(),
        })

    missing = "Chưa có dữ liệu" if vi else "Missing data"
    unknown = "Chưa rõ" if vi else "TBD"
    blocks = [
        f"# {title}",
        f"## {'Thành phần tham gia' if vi else 'Attendees'}\n" + ("\n".join(f"- {x}" for x in attendees) if attendees else f"- {missing}"),
        f"## {'Nội dung trao đổi chính' if vi else 'Key Discussion'}\n" + ("\n".join(f"- {x}" for x in key_points) if key_points else f"- {missing}"),
        f"## {'Quyết định quan trọng' if vi else 'Key Decisions'}\n" + ("\n".join(f"- {x}" for x in decisions) if decisions else f"- {missing}"),
    ]
    if action_rows:
        action_lines = []
        for idx, row in enumerate(action_rows, start=1):
            action_lines.append(
                f"{idx}. **What:** {row['task']}\n"
                f"   **Who:** {row['assignee'] or unknown}\n"
                f"   **When:** {row['deadline'] or unknown}"
            )
        blocks.append(f"## {'Action items (What - Who - When)' if vi else 'Action Items (What - Who - When)'}\n" + "\n".join(action_lines))
    else:
        blocks.append(f"## {'Action items (What - Who - When)' if vi else 'Action Items (What - Who - When)'}\n- {missing}")

    if summary:
        blocks.append(f"## {'Tóm tắt' if vi else 'Summary'}\n{summary}")
    return "\n\n".join(blocks).strip()


def _normalize_minutes_markdown(raw: str, lang: str = "vi") -> str:
    text = (raw or "").strip()
    if not text:
        return ""
    legacy = _parse_legacy_minutes_payload(text)
    if legacy:
        return _legacy_minutes_to_markdown(legacy, lang)
    return text


def _markdown_to_docx(markdown_text: str, out_path: Path):
    from docx import Document
    import re

    def _strip_inline(md_line: str) -> str:
        line = re.sub(r"`([^`]+)`", r"\1", md_line)
        line = re.sub(r"\*\*([^*]+)\*\*", r"\1", line)
        line = re.sub(r"\*([^*]+)\*", r"\1", line)
        return line.strip()

    doc = Document()
    for raw in markdown_text.replace("\r", "").split("\n"):
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            continue
        h = re.match(r"^(#{1,3})\s+(.+)$", line)
        if h:
            level = min(3, len(h.group(1)))
            doc.add_heading(_strip_inline(h.group(2)), level=level)
            continue
        ul = re.match(r"^[-*]\s+(.+)$", line)
        if ul:
            doc.add_paragraph(_strip_inline(ul.group(1)), style="List Bullet")
            continue
        ol = re.match(r"^\d+\.\s+(.+)$", line)
        if ol:
            doc.add_paragraph(_strip_inline(ol.group(1)), style="List Number")
            continue
        doc.add_paragraph(_strip_inline(line))
    doc.save(out_path)


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup / shutdown."""
    db.init()
    lang = db.get_setting("app_language") or "vi"
    print(t("starting", lang))

    # ── Diarizer model init (BLOCKING — app won't serve until loaded) ──
    print("[main] Initializing diarizer model (blocking)...")
    try:
        diarizer._init_model()
        if not diarizer._session:
            raise RuntimeError("Model loaded but instance is None")
        print(f"[main] ✓ Diarizer model ready (model_loaded={diarizer._model_loaded})")
    except Exception as e:
        print(f"[main] ✗ Diarizer model FAILED: {e}")
        print("[main] Exiting — app cannot run without diarization model")
        import sys
        sys.exit(1)

    # ── Riva warmup (background — not critical for app start) ──
    import threading

    def _warmup_riva():
        try:
            nvidia_key = db.get_setting("nvidia_api_key") or os.environ.get("NVIDIA_API_KEY", "")
            if nvidia_key:
                from stt import _get_riva_asr, get_nvidia_model
                stt_lang_warmup = db.get_setting("stt_language") or "vi"
                lang_code = get_language_code(stt_lang_warmup)
                model = get_nvidia_model(lang_code)
                _get_riva_asr(nvidia_key, model["function_id"])
                print(t("riva_connected", lang))
        except Exception as e:
            print(f"{t('riva_warmup_fail', lang)}: {e}")

    threading.Thread(target=_warmup_riva, daemon=True).start()

    yield
    print(t("shutting_down", lang))


app = FastAPI(title="VoiceScribe Sidecar", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ─── Health ───
@app.get("/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}


# ─── Live Log Stream (SSE) ───
@app.get("/logs")
async def logs_stream():
    async def event_gen():
        q: _queue.Queue = _queue.Queue(maxsize=200)
        with _log_lock:
            _log_subscribers.append(q)
            for line in _log_buffer:
                yield f"data: {line}\n\n"
        try:
            while True:
                try:
                    line = q.get_nowait()
                    yield f"data: {line}\n\n"
                except _queue.Empty:
                    await asyncio.sleep(0.3)
        except asyncio.CancelledError:
            pass
        finally:
            with _log_lock:
                if q in _log_subscribers:
                    _log_subscribers.remove(q)

    return StreamingResponse(event_gen(), media_type="text/event-stream")


@app.get("/diarizer-status")
async def diarizer_status():
    """Diagnostic endpoint: check if model loaded and config state."""
    return {
        "model_loaded": diarizer._model_loaded,
        "model_ok": diarizer._session is not None,
        "source": diarizer._source,
        "profile_count": len(diarizer._profiles),
        "config": {
            "match_threshold": diarizer.cfg("match_threshold"),
            "pitch_penalty_factor": diarizer.cfg("pitch_penalty_factor"),
            "switch_confirm_hits": diarizer.cfg("switch_confirm_hits"),
            "same_zone_pitch_diff_male": diarizer.cfg("same_zone_weak_pitch_diff_male"),
            "same_zone_pitch_diff_female": diarizer.cfg("same_zone_weak_pitch_diff_female"),
        },
    }


# ─── Transcribe + Diarize (combined, parallel) ───
@app.post("/transcribe-diarize")
async def transcribe_diarize(audio: UploadFile = File(...)):
    """Process audio chunk: transcribe and assign one stable speaker per chunk."""
    content = await audio.read()

    # Detect format from content: WAV starts with "RIFF", else assume webm
    suffix = ".wav" if content[:4] == b"RIFF" else ".webm"
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tmp_path = tmp.name
    try:
        tmp.write(content)
        tmp.flush()
        tmp.close()

        # Run transcription and diarization in parallel
        loop = asyncio.get_event_loop()
        stt_lang = db.get_setting("stt_language") or "vi"
        nvidia_key = db.get_setting("nvidia_api_key") or os.environ.get("NVIDIA_API_KEY", "")
        riva_lang = get_language_code(stt_lang)
        transcribe_task = loop.run_in_executor(None, transcribe_nvidia, tmp_path, nvidia_key, riva_lang)
        diarize_task = loop.run_in_executor(None, diarizer.identify_speaker, tmp_path)

        results = await asyncio.gather(transcribe_task, diarize_task, return_exceptions=True)

        # Handle STT result
        text = ""
        if isinstance(results[0], Exception):
            print(f"[transcribe-diarize] STT error: {results[0]}")
        else:
            text = results[0] or ""

        # Handle diarization result
        if isinstance(results[1], Exception):
            print(f"[transcribe-diarize] diarize error: {results[1]}")
            speaker_info = {"speaker": "Speaker 1", "speaker_id": 0}
        else:
            speaker_info = results[1]

        # Filter hallucinations
        text = filter_hallucinations(text)
        if not text.strip():
            return {"text": "", "segments": [], "speakers": len(diarizer._profiles)}

        speaker = speaker_info.get("speaker", "Speaker 1")
        speaker_id = speaker_info.get("speaker_id", 0)
        chunk_id = f"chunk-{int(time.time() * 1000)}-{uuid4().hex[:8]}"
        text = text.strip()

        return {
            "text": text,
            "chunk_id": chunk_id,
            "segments": [{
                "speaker": speaker,
                "speaker_id": speaker_id,
                "chunk_id": chunk_id,
                "text": text,
            }],
            "speakers": len(diarizer._profiles),
        }
    finally:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass


# ─── Nvidia Streaming WebSocket ───
@app.websocket("/ws/nvidia-stream")
async def nvidia_stream_ws(websocket: WebSocket):
    """WebSocket endpoint for real-time Nvidia Riva streaming STT.

    Client sends: binary frames (raw PCM 16kHz mono int16 LE audio)
    Client receives:
      - transcript frames: {"text": "...", "is_final": true/false, "speaker": "...", "speaker_id": 0, "chunk_id": "...?"}
    Client sends text "STOP" to end the stream.
    Query params:
      - source=system  → use system audio diarization config (lower thresholds)
      - source=web     → use default browser audio config (default)
      - meeting_id=123 → optional, append raw PCM stream to meeting audio archive
    """
    await websocket.accept()

    # Detect audio source and apply diarization config
    source = websocket.query_params.get("source", "web")
    if source == "system":
        diarizer.set_source("system")
    else:
        diarizer.set_source("web")

    meeting_id_raw = websocket.query_params.get("meeting_id")
    archive_fh = None
    if meeting_id_raw:
        try:
            meeting_id = int(meeting_id_raw)
            meeting = db.get_meeting(meeting_id)
            if meeting:
                audio_dir = _voicescribe_data_dir() / "audio"
                audio_dir.mkdir(parents=True, exist_ok=True)
                archive_path = audio_dir / f"meeting_{meeting_id}.pcm"
                archive_fh = archive_path.open("ab")
                # Persist archive path so /meetings/{id}/audio can resolve it
                db.update_meeting(meeting_id, audio_path=str(archive_path))
        except Exception as e:
            print(f"[ws:nvidia-stream] archive setup failed: {e}")

    nvidia_key = db.get_setting("nvidia_api_key") or os.environ.get("NVIDIA_API_KEY", "")
    stt_lang = db.get_setting("stt_language") or "vi"
    riva_lang = get_language_code(stt_lang)

    if not nvidia_key:
        await websocket.send_json({"error": "NVIDIA_API_KEY not set"})
        await websocket.close()
        return

    streamer = NvidiaStreamingSTT(nvidia_key, riva_lang)
    loop = asyncio.get_event_loop()

    # Start streaming in a thread
    try:
        await loop.run_in_executor(None, streamer.start)
    except Exception as e:
        await websocket.send_json({"error": str(e)})
        await websocket.close()
        return

    # Bridge blocking results() iterator to async via thread + asyncio.Queue
    result_queue = asyncio.Queue()
    audio_buffer = bytearray()    # Rolling buffer for interim diarization
    segment_buffer = bytearray()  # Segment buffer since last final result
    interim_speaker = {"speaker": "Speaker 1", "speaker_id": 0}
    last_interim_diarize_at = 0.0
    interim_diarize_task = None
    bytes_per_second = 16000 * 2  # 16kHz mono int16
    # Diarization windows — same proven values for all models.
    # The diarizer needs enough audio to reliably distinguish speakers.
    is_vietnamese = riva_lang.startswith("vi")
    interim_window_sec = float(os.getenv("NVIDIA_DIARIZE_INTERIM_WINDOW_SEC", "10.0"))
    final_window_sec = float(os.getenv("NVIDIA_DIARIZE_FINAL_WINDOW_SEC", "14.0"))
    final_recent_window_sec = float(os.getenv("NVIDIA_DIARIZE_FINAL_RECENT_WINDOW_SEC", "4.0"))
    context_keep_sec = float(os.getenv("NVIDIA_DIARIZE_CONTEXT_KEEP_SEC", "1.5"))
    boundary_fix_window_sec = float(os.getenv("NVIDIA_DIARIZE_BOUNDARY_FIX_WINDOW_SEC", "2.4"))
    boundary_fix_max_words = int(os.getenv("NVIDIA_DIARIZE_BOUNDARY_FIX_MAX_WORDS", "8"))
    # Multilingual RNNT emits fewer finals → interim diarization runs more often
    interim_diarize_interval = float(os.getenv("NVIDIA_DIARIZE_INTERIM_INTERVAL",
                                               "1.2" if is_vietnamese else "0.8"))
    last_final_chunk_id = ""
    last_final_speaker_id = None
    last_final_text = ""
    last_final_at = 0.0

    def _tail_bytes(buf: bytes | bytearray, seconds: float) -> bytes:
        max_bytes = max(1, int(seconds * bytes_per_second))
        if len(buf) <= max_bytes:
            return bytes(buf)
        return bytes(buf[-max_bytes:])

    def _read_results():
        for result in streamer.results():
            asyncio.run_coroutine_threadsafe(result_queue.put(result), loop)
        asyncio.run_coroutine_threadsafe(result_queue.put(None), loop)

    import threading
    import wave
    result_thread = threading.Thread(target=_read_results, daemon=True)
    result_thread.start()

    def _diarize_buffer(buf: bytes, update_profiles: bool = True) -> dict:
        """Write buffered PCM to temp WAV and run diarization."""
        try:
            duration_s = len(buf) / (16000 * 2)  # 16kHz, 16-bit = 2 bytes/sample
            if duration_s < 0.5:
                return {"speaker": "Speaker 1", "speaker_id": 0}
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
            with wave.open(tmp.name, 'wb') as wf:
                wf.setnchannels(1)
                wf.setsampwidth(2)  # 16-bit
                wf.setframerate(16000)
                wf.writeframes(buf)
            print(f"[ws:diarize] buffer={duration_s:.1f}s update={update_profiles} source={source}")
            speaker_info = diarizer.identify_speaker(tmp.name, update_profiles=update_profiles)
            os.unlink(tmp.name)
            return speaker_info
        except Exception as e:
            print(f"[ws:diarize] error: {e}")
            return {"speaker": "Speaker 1", "speaker_id": 0}

    async def _refresh_interim_speaker(buf_copy: bytes):
        nonlocal interim_speaker, segment_buffer
        try:
            speaker_info = await loop.run_in_executor(None, _diarize_buffer, buf_copy, False)
            new_speaker_id = speaker_info.get("speaker_id", 0)
            old_speaker_id = interim_speaker.get("speaker_id", 0)

            interim_speaker = {
                "speaker": speaker_info.get("speaker", "Speaker 1"),
                "speaker_id": new_speaker_id,
            }

            # If speaker changed, tell frontend to split the transcript immediately
            if new_speaker_id != old_speaker_id:
                print(f"[ws:interim-diarize] Speaker change: S{old_speaker_id+1} → S{new_speaker_id+1}")
                # Reset segment buffer so next final only covers the new speaker's audio
                segment_buffer = bytearray(_tail_bytes(audio_buffer, context_keep_sec))
                try:
                    await websocket.send_json({
                        "type": "speaker_split",
                        "speaker": interim_speaker["speaker"],
                        "speaker_id": new_speaker_id,
                    })
                except Exception:
                    pass
        except Exception as e:
            print(f"[ws:interim-diarize] error: {e}")

    # Consumer from asyncio queue
    async def _send_results():
        nonlocal audio_buffer, segment_buffer, interim_speaker, last_interim_diarize_at, interim_diarize_task
        nonlocal last_final_chunk_id, last_final_speaker_id, last_final_text, last_final_at
        while True:
            result = await result_queue.get()
            if result is None:
                break
            try:
                msg = {
                    "text": result["text"],
                    "is_final": result["is_final"],
                    "speaker": "Speaker 1",
                    "speaker_id": 0,
                }

                if not result["is_final"]:
                    # Interim result: send immediately for live typing
                    msg["speaker"] = interim_speaker.get("speaker", "Speaker 1")
                    msg["speaker_id"] = interim_speaker.get("speaker_id", 0)

                    # Try to refresh speaker guess in background every ~1.2s
                    # (enough audio for diarization, but without blocking interim text flow)
                    now = time.time()
                    if len(audio_buffer) >= 16000 * 2 and (now - last_interim_diarize_at) >= interim_diarize_interval:
                        if interim_diarize_task is None or interim_diarize_task.done():
                            last_interim_diarize_at = now
                            interim_diarize_task = asyncio.create_task(
                                _refresh_interim_speaker(_tail_bytes(audio_buffer, interim_window_sec))
                            )

                    await websocket.send_json(msg)
                else:
                    # Final result: run diarization first, then send with correct speaker
                    if result["text"] and len(segment_buffer) > 0:
                        buf_copy = _tail_bytes(segment_buffer, final_window_sec)
                        recent_buf = _tail_bytes(segment_buffer, final_recent_window_sec)
                        if len(recent_buf) < 16000:
                            recent_buf = buf_copy
                        # Start next segment with a little context for smoother boundary
                        segment_buffer = bytearray(_tail_bytes(audio_buffer, context_keep_sec))
                        speaker_info = await loop.run_in_executor(None, _diarize_buffer, recent_buf)
                        msg["speaker"] = speaker_info.get("speaker", "Speaker 1")
                        msg["speaker_id"] = speaker_info.get("speaker_id", 0)
                        chunk_id = f"chunk-{int(time.time() * 1000)}-{uuid4().hex[:8]}"
                        msg["chunk_id"] = chunk_id
                        interim_speaker = {"speaker": msg["speaker"], "speaker_id": msg["speaker_id"]}

                        # Boundary correction: if switch is detected slightly late, relabel
                        # a very short previous chunk to current speaker for better readability.
                        now_ts = time.time()
                        word_count = len(last_final_text.strip().split()) if last_final_text else 0
                        if (
                            last_final_chunk_id
                            and last_final_speaker_id is not None
                            and last_final_speaker_id != msg["speaker_id"]
                            and (now_ts - last_final_at) <= boundary_fix_window_sec
                            and 0 < word_count <= boundary_fix_max_words
                        ):
                            await websocket.send_json({
                                "type": "speaker_correction",
                                "chunk_id": last_final_chunk_id,
                                "speaker": msg["speaker"],
                                "speaker_id": msg["speaker_id"],
                            })

                    else:
                        msg["speaker"] = interim_speaker.get("speaker", "Speaker 1")
                        msg["speaker_id"] = interim_speaker.get("speaker_id", 0)
                        if result["text"]:
                            chunk_id = f"chunk-{int(time.time() * 1000)}-{uuid4().hex[:8]}"
                            msg["chunk_id"] = chunk_id

                    if msg.get("chunk_id"):
                        last_final_chunk_id = msg["chunk_id"]
                        last_final_speaker_id = msg.get("speaker_id", 0)
                        last_final_text = result.get("text", "") or ""
                        last_final_at = time.time()
                    await websocket.send_json(msg)

            except WebSocketDisconnect:
                break

    send_task = asyncio.create_task(_send_results())

    # Producer: receive audio from client
    try:
        while True:
            data = await websocket.receive()
            if data.get("type") == "websocket.disconnect":
                break
            if "bytes" in data:
                audio_bytes = data["bytes"]
                streamer.feed_audio(audio_bytes)
                audio_buffer.extend(audio_bytes)  # Buffer for diarization
                segment_buffer.extend(audio_bytes)
                if archive_fh is not None:
                    try:
                        archive_fh.write(audio_bytes)
                    except Exception as e:
                        print(f"[ws:nvidia-stream] archive write failed: {e}")
                # Prevent unbounded growth if ASR finalization is delayed
                max_interim_bytes = max(1, int(interim_window_sec * bytes_per_second))
                max_segment_bytes = max(1, int(final_window_sec * bytes_per_second))
                if len(audio_buffer) > max_interim_bytes:
                    audio_buffer = bytearray(audio_buffer[-max_interim_bytes:])
                if len(segment_buffer) > max_segment_bytes:
                    segment_buffer = bytearray(segment_buffer[-max_segment_bytes:])
            elif "text" in data:
                if data["text"] == "STOP":
                    break
    except WebSocketDisconnect:
        pass
    finally:
        streamer.stop()
        result_thread.join(timeout=5)
        if interim_diarize_task and not interim_diarize_task.done():
            interim_diarize_task.cancel()
        send_task.cancel()
        if archive_fh is not None:
            try:
                archive_fh.close()
            except Exception:
                pass
        try:
            await websocket.close()
        except Exception:
            pass


# ─── Diarize Reset ───
@app.post("/diarize-reset")
async def diarize_reset():
    diarizer.reset()
    return {"ok": True}


# ─── Translation (SSE streaming) ───
@app.post("/translate")
async def translate(request: Request):
    body = await request.json()
    text = body.get("text", "")
    target_lang = body.get("targetLang", "en")

    if not text:
        return JSONResponse({"error": "No text"}, status_code=400)

    from translate import translate_stream
    return StreamingResponse(
        translate_stream(text, target_lang, db),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


# ─── Summarization (SSE streaming) ───
@app.post("/summarize")
async def summarize(request: Request):
    body = await request.json()
    meeting_id = body.get("meetingId")
    language = body.get("language", "vi")
    transcript = body.get("transcript")

    # If transcript provided directly (unsaved meeting), use it
    if not transcript:
        if not meeting_id:
            return JSONResponse({"error": "No meetingId or transcript provided"}, status_code=400)
        meeting = db.get_meeting(meeting_id)
        if not meeting:
            return JSONResponse({"error": "Meeting not found"}, status_code=404)
        transcript = meeting["transcript"]

    from summarize import summarize_stream
    return StreamingResponse(
        summarize_stream(transcript, language, db),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive"},
    )


# ─── Settings ───
@app.get("/settings")
async def get_settings():
    return db.get_all_settings()


@app.post("/settings")
async def save_settings(request: Request):
    body = await request.json()
    old_lang = db.get_setting("stt_language") or ""
    for key, value in body.items():
        db.set_setting(key, str(value))
    # Reset ASR cache if language changed (different function_id needed)
    new_lang = body.get("stt_language", old_lang)
    if new_lang != old_lang or "nvidia_api_key" in body:
        try:
            from stt import _reset_riva_asr
            _reset_riva_asr()  # Clear all cached ASR services
            print(f"[settings] ASR cache cleared (language: {old_lang} → {new_lang})")
        except Exception as e:
            print(f"[settings] ASR cache reset failed: {e}")
    return {"ok": True}


# ─── Meetings CRUD ───
@app.get("/meetings")
async def list_meetings():
    return db.get_all_meetings()


@app.post("/meetings")
async def create_meeting(request: Request):
    body = await request.json()
    mid = db.create_meeting(
        title=body.get("title", "Untitled"),
        transcript=json.dumps(body.get("transcript", [])) if isinstance(body.get("transcript"), list) else body.get(
            "transcript", ""),
        summary=body.get("summary", ""),
        audio_duration=body.get("audioDuration", 0),
        language=body.get("language", "vi"),
    )
    return {"id": mid}


@app.get("/meetings/{meeting_id}")
async def get_meeting(meeting_id: int):
    m = db.get_meeting(meeting_id)
    if not m:
        return JSONResponse({"error": "Not found"}, status_code=404)
    return m


@app.get("/meetings/{meeting_id}/audio")
async def download_meeting_audio(meeting_id: int, format: str = "wav"):
    m = db.get_meeting(meeting_id)
    audio_path = (m.get("audio_path") or "") if m else ""
    source_path = Path(audio_path) if audio_path else None
    if source_path is None or not source_path.exists() or not source_path.is_file():
        # Fallback: search audio files by meeting ID pattern
        audio_dir = _voicescribe_data_dir() / "audio"
        for ext in (".wav", ".mp4", ".m4a", ".mp3", ".webm", ".pcm"):
            candidate = audio_dir / f"meeting_{meeting_id}{ext}"
            if candidate.exists() and candidate.is_file():
                source_path = candidate
                if m:
                    try:
                        db.update_meeting(meeting_id, audio_path=str(candidate))
                    except Exception:
                        pass
                break
    if source_path is None:
        return JSONResponse({"error": "Audio not found"}, status_code=404)
    if not source_path.exists() or not source_path.is_file():
        return JSONResponse({"error": "Audio not found"}, status_code=404)

    export_format = (format or "wav").strip().lower()
    if export_format not in {"wav", "mp4"}:
        return JSONResponse({"error": "format must be wav or mp4"}, status_code=400)

    export_path = source_path
    cleanup_task = None
    if source_path.suffix.lower() != f".{export_format}":
        # Fast path: PCM → WAV using Python wave module (no ffmpeg, no blocking)
        if source_path.suffix.lower() == ".pcm" and export_format == "wav":
            import wave as _wave
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".wav")
            tmp_path = Path(tmp.name)
            tmp.close()
            try:
                pcm_data = source_path.read_bytes()
                with _wave.open(str(tmp_path), 'wb') as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(2)  # 16-bit
                    wf.setframerate(16000)
                    wf.writeframes(pcm_data)
                export_path = tmp_path
                cleanup_task = BackgroundTask(_safe_unlink, str(export_path))
            except Exception as e:
                _safe_unlink(str(tmp_path))
                return JSONResponse(
                    {"error": f"PCM to WAV conversion failed: {e}"},
                    status_code=500,
                )
        else:
            try:
                import asyncio
                export_path = await asyncio.to_thread(_transcode_audio_for_export, source_path, export_format)
                cleanup_task = BackgroundTask(_safe_unlink, str(export_path))
            except Exception as e:
                return JSONResponse(
                    {"error": f"Audio convert to {export_format} failed: {e}"},
                    status_code=500,
                )

    media_type = _audio_media_type(export_path.suffix)
    return FileResponse(
        str(export_path),
        media_type=media_type,
        filename=f"meeting_{meeting_id}.{export_format}",
        background=cleanup_task,
    )


@app.get("/meetings/{meeting_id}/minutes")
async def download_meeting_minutes(meeting_id: int, format: str = "md"):
    m = db.get_meeting(meeting_id)
    if not m:
        return JSONResponse({"error": "Not found"}, status_code=404)

    summary = str(m.get("summary") or "").strip()
    if not summary:
        return JSONResponse({"error": "Minutes not found"}, status_code=404)

    language = str(m.get("language") or "vi")
    markdown = _normalize_minutes_markdown(summary, language)
    if not markdown:
        return JSONResponse({"error": "Minutes not found"}, status_code=404)

    export_format = (format or "md").strip().lower()
    if export_format not in {"md", "docx"}:
        return JSONResponse({"error": "format must be md or docx"}, status_code=400)

    raw_title = str(m.get("title") or f"meeting-{meeting_id}").strip()
    safe_title = "".join(ch for ch in raw_title if ch not in '\\/:*?"<>|').strip() or f"meeting-{meeting_id}"

    if export_format == "md":
        return Response(
            content=markdown.encode("utf-8"),
            media_type="text/markdown; charset=utf-8",
            headers={"Content-Disposition": f'attachment; filename="{safe_title}-minutes.md"'},
        )

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_path = Path(tmp.name)
    tmp.close()
    try:
        _markdown_to_docx(markdown, tmp_path)
    except ModuleNotFoundError:
        _safe_unlink(str(tmp_path))
        return JSONResponse({"error": "python-docx is not installed"}, status_code=500)
    except Exception as e:
        _safe_unlink(str(tmp_path))
        return JSONResponse({"error": f"docx export failed: {e}"}, status_code=500)

    return FileResponse(
        str(tmp_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=f"{safe_title}-minutes.docx",
        background=BackgroundTask(_safe_unlink, str(tmp_path)),
    )


@app.put("/meetings/{meeting_id}")
async def update_meeting(meeting_id: int, request: Request):
    body = await request.json()
    m = db.get_meeting(meeting_id)
    if not m:
        return JSONResponse({"error": "Not found"}, status_code=404)
    # Map camelCase from frontend to snake_case for DB
    updates = {}
    key_map = {
        "title": "title", "transcript": "transcript", "summary": "summary",
        "translations": "translations", "audioDuration": "audio_duration",
        "language": "language", "status": "status", "audioPath": "audio_path",
    }
    for js_key, db_key in key_map.items():
        if js_key in body:
            val = body[js_key]
            if isinstance(val, list):
                val = json.dumps(val)
            updates[db_key] = val
    if updates:
        db.update_meeting(meeting_id, **updates)
    return {"ok": True}


@app.delete("/meetings/{meeting_id}")
async def delete_meeting(meeting_id: int):
    db.delete_meeting(meeting_id)
    return {"ok": True}


# ─── Drafts ───
@app.post("/drafts")
async def create_draft(request: Request):
    body = await request.json()
    mid = db.create_meeting(
        title=body.get("title", "Draft"),
        transcript="",
        summary="",
        audio_duration=0,
        language="vi",
        status="draft",
    )
    return {"id": mid}


@app.patch("/drafts/{draft_id}")
async def append_draft(draft_id: int, request: Request):
    body = await request.json()
    duration = body.get("audioDuration", 0)
    m = db.get_meeting(draft_id)
    if not m:
        return JSONResponse({"error": "Draft not found"}, status_code=404)

    # Support both: structured part (JSON) or plain text
    part = body.get("part")  # {text, speaker, speakerId, startTime, endTime}
    append_text = body.get("appendText", "")

    current = m.get("transcript", "") or ""
    if part:
        # Store as JSON array of parts
        import json
        try:
            parts = json.loads(current) if current.startswith("[") else []
        except Exception:
            parts = []

        # Upsert final snapshot instead of always append:
        # when same speaker + same startTime, replace tail to avoid duplicate-growing rows.
        replaced = False
        if parts and isinstance(parts[-1], dict):
            last = parts[-1]
            same_speaker = (
                str(last.get("speakerId", "")) == str(part.get("speakerId", ""))
                or str(last.get("speaker", "")) == str(part.get("speaker", ""))
            )
            same_start = str(last.get("startTime", "")) == str(part.get("startTime", ""))
            if same_speaker and same_start:
                merged = dict(last)
                merged.update(part)
                # keep non-empty accumulated chunkIds if present
                if isinstance(last.get("chunkIds"), list) and isinstance(part.get("chunkIds"), list):
                    ids = []
                    for cid in [*last.get("chunkIds", []), *part.get("chunkIds", [])]:
                        if isinstance(cid, str) and cid and cid not in ids:
                            ids.append(cid)
                    if ids:
                        merged["chunkIds"] = ids
                parts[-1] = merged
                replaced = True

        if not replaced:
            parts.append(part)
        db.update_meeting(draft_id, transcript=json.dumps(parts, ensure_ascii=False), audio_duration=duration)
    elif append_text:
        db.update_meeting(draft_id, transcript=current + "\n" + append_text, audio_duration=duration)
    return {"ok": True}


@app.patch("/drafts/{draft_id}/audio")
async def append_draft_audio(draft_id: int, audio: UploadFile = File(...)):
    m = db.get_meeting(draft_id)
    if not m:
        return JSONResponse({"error": "Draft not found"}, status_code=404)

    payload = await audio.read()
    if not payload:
        return {"ok": True, "bytes": 0}

    audio_path = (m.get("audio_path") or "").strip()
    if audio_path:
        target = Path(audio_path)
        target.parent.mkdir(parents=True, exist_ok=True)
    else:
        audio_dir = _voicescribe_data_dir() / "audio"
        audio_dir.mkdir(parents=True, exist_ok=True)
        suffix = Path(audio.filename or "").suffix.lower() or ".webm"
        target = audio_dir / f"meeting_{draft_id}{suffix}"
        db.update_meeting(draft_id, audio_path=str(target))

    with target.open("ab") as f:
        f.write(payload)

    return {"ok": True, "bytes": len(payload), "audioPath": str(target)}


# ─── Diagnostics ───
@app.get("/diagnose")
async def diagnose(lang: str = "vi"):
    import httpx
    results = {"stt": {"status": "unknown", "message": ""}, "llm": {"status": "unknown", "message": ""}}
    loop = asyncio.get_event_loop()

    # Check STT (Nvidia Riva)
    nvidia_key = db.get_setting("nvidia_api_key") or os.getenv("NVIDIA_API_KEY", "")
    if not nvidia_key:
        results["stt"] = {"status": "warning", "message": t("nvidia_key_missing", lang)}
    else:
        try:
            from stt import _get_riva_asr, get_nvidia_model
            stt_lang_diag = db.get_setting("stt_language") or "vi"
            lang_code = get_language_code(stt_lang_diag)
            model = get_nvidia_model(lang_code)
            await loop.run_in_executor(None, _get_riva_asr, nvidia_key, model["function_id"])
            results["stt"] = {"status": "ok", "message": t("nvidia_connected", lang)}
        except Exception as e:
            results["stt"] = {"status": "error", "message": f"{t('nvidia_connect_fail', lang)}: {str(e)[:80]}"}

    # Check LLM
    llm_key = db.get_setting("llm_api_key") or os.getenv("LLM_API_KEY", "")
    llm_url = db.get_setting("llm_base_url") or os.getenv("LLM_BASE_URL", "")
    if not llm_key:
        results["llm"] = {"status": "warning", "message": t("llm_key_missing", lang)}
    else:
        try:
            base = llm_url.rstrip("/") if llm_url else "https://api.openai.com/v1"
            async with httpx.AsyncClient(timeout=8) as client:
                r = await client.get(f"{base}/models", headers={"Authorization": f"Bearer {llm_key}"})
                if r.status_code == 200:
                    results["llm"] = {"status": "ok", "message": t("llm_connected", lang)}
                else:
                    results["llm"] = {"status": "error", "message": t("llm_key_invalid", lang)}
        except Exception:
            results["llm"] = {"status": "error", "message": t("llm_connect_fail", lang)}

    results["backend"] = "nvidia"
    return results


def filter_hallucinations(text: str) -> str:
    import re
    for pattern in HALLUCINATION_PATTERNS:
        if re.search(pattern, text, re.IGNORECASE):
            return ""
    if len(text.strip()) <= 3:
        return ""
    return text


if __name__ == "__main__":
    import traceback

    # ── Crash-safe logging ──
    def _crash_log_path():
        data_dir = os.getenv("VOICESCRIBE_DATA", os.path.join(os.path.expanduser("~"), ".voicescribe"))
        os.makedirs(data_dir, exist_ok=True)
        return os.path.join(data_dir, "sidecar-crash.log")

    try:
        # Log startup environment for debugging
        log_path = _crash_log_path()
        with open(log_path, "w") as f:
            f.write(f"[startup] frozen={getattr(sys, 'frozen', False)}\n")
            f.write(f"[startup] _MEIPASS={getattr(sys, '_MEIPASS', 'N/A')}\n")
            f.write(f"[startup] executable={sys.executable}\n")
            f.write(f"[startup] cwd={os.getcwd()}\n")
            f.write(f"[startup] platform={sys.platform}\n")
            f.write(f"[startup] Starting uvicorn...\n")

        port = int(os.getenv("SIDECAR_PORT", "8765"))
        uvicorn.run(app, host="127.0.0.1", port=port, log_level="info")
    except Exception as e:
        tb = traceback.format_exc()
        print(f"[CRASH] {e}\n{tb}", file=sys.stderr)
        try:
            with open(_crash_log_path(), "a") as f:
                f.write(f"\n[CRASH] {e}\n{tb}\n")
        except Exception:
            pass
        sys.exit(1)

